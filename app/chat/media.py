import base64
from io import BytesIO
from fastapi import UploadFile
from openai import AsyncOpenAI
from app.core.config import get_settings
from pypdf import PdfReader
from docx import Document

settings = get_settings()

def extract_doc_text(data:bytes)->str:
    doc = Document(BytesIO(data))
    parts=[p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells=[c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("|".join(cells))
    return "\n".join(parts)

def extract_pdf_text(data:bytes,max_pages:int=50)->str:
    reader = PdfReader(BytesIO(data))
    parts:list[str]=[]
    for i,page in enumerate(reader.pages):
        if i>max_pages:
            break
        parts.append(page.extract_text() or "")
    text="\n\n".join(parts).strip()
    if len(text)<100 and len(reader.pages)>=5:
        return "[это скан, OCR пока не поддкрживается]"
    return text

async def whisper_transcribe(audio_bytes:bytes,filename:str,llm_service) -> str:
    f=BytesIO(audio_bytes)
    f.name=filename
    result=await llm_service.audio.transcriptions.create(
        model="stt-openai/gpt-4o-mini-transcribe",
        file=f,
    )
    return result.text

async def media_to_part(media:UploadFile,llm_client)->dict:
    mime=media.content_type
    data=await media.read()
    if mime.startswith("audio/") or mime=="application/ogg":
        transcript=await whisper_transcribe(data,media.filename
                                            or "audio.ogg",llm_client)
        return {
            "type":"text",
            "text":f"Пользователь сказал голосом:\n {transcript}"
        }

    if mime.startswith("image/"):
        b64=base64.b64encode(data).decode()
        return {
            "type":"image_url",
            "image_url":{"url":f"data:{mime};base64,{b64}"}
        }

    if mime.startswith("wordprocessingml.document"):
        return {
            "type": "text",
            "text": f"[Документ DOCX]:\n{extract_doc_text(data)[:30000]}"
        }

    if mime=="application/pdf":
        return {
            "type": "text",
            "text": f"[Документ PDF]:\n{extract_pdf_text(data)[:30000]}"
        }

    raise ValueError(f"Unsupported media type: {mime}")